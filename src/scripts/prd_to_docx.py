#!/usr/bin/env python3
"""prd_to_docx.py — 本地 Word 导出（蒸馏 A3 pm-prd-workflow generate_prd_docx.py）

把 Markdown PRD（或任意产物）转换为 .docx，供本地存档 / 非飞书交付 / 评审打印。
不依赖 lark-cli（飞书发布走 prd_publish.py）；本脚本是纯本地 Offline 能力。

用法:
  python3 src/scripts/prd_to_docx.py <input.md> [--output out.docx] [--title 标题]

依赖:
  python-docx（未安装时给出提示；可选安装：pip install python-docx）

兼容性:
  - 解析 Markdown：标题(#/##/###)、表格、代码块(fenced)、列表(-/1.)、粗体、链接
  - 表格按 Markdown 管道分隔解析，行/列保留
  - 不支持：内嵌图片、复杂嵌套列表、HTML 块（原样保留为文本）
  - 产物 frontmatter（--- 区间）跳过不导出

失败模式:
  - python-docx 未安装 → stderr 提示，退出码 2（不静默降级）
"""

from __future__ import annotations

import argparse
import re
import sys
from pathlib import Path


def _strip_frontmatter(text: str) -> str:
    """Remove leading YAML frontmatter (--- ... ---) if present."""
    m = re.match(r"\A(?:<!--.*?-->\s*)?---\s*\n.*?\n---\s*\n", text, re.DOTALL)
    return text[m.end():] if m else text


def _iter_blocks(text: str):
    """Yield (kind, content) blocks from a Markdown document.

    kind ∈ {h1,h2,h3,code,table,list,para}
    """
    lines = text.splitlines()
    i = 0
    n = len(lines)
    while i < n:
        line = lines[i]
        stripped = line.strip()
        if not stripped:
            i += 1
            continue
        # fenced code block
        m = re.match(r"^```(\w*)\s*$", stripped)
        if m:
            lang = m.group(1)
            buf = [line]
            i += 1
            while i < n and not lines[i].strip().startswith("```"):
                buf.append(lines[i])
                i += 1
            if i < n:
                buf.append(lines[i])
                i += 1
            yield ("code", "\n".join(buf))
            continue
        # headings
        hm = re.match(r"^(#{1,6})\s+(.+?)\s*$", stripped)
        if hm:
            level = len(hm.group(1))
            yield (f"h{min(level, 3)}", hm.group(2))
            i += 1
            continue
        # table: current line starts with '|' and next is separator row
        if stripped.startswith("|") and i + 1 < n:
            nxt = lines[i + 1].strip()
            if re.match(r"^\|[\s:\-|]+\|$", nxt):
                rows = []
                rows.append(lines[i])
                rows.append(nxt)
                i += 2
                while i < n and lines[i].strip().startswith("|"):
                    rows.append(lines[i])
                    i += 1
                yield ("table", "\n".join(rows))
                continue
        # list item
        if re.match(r"^(\s*[-*+]|\s*\d+\.)\s+", line):
            buf = [line]
            i += 1
            while i < n and re.match(r"^(\s*[-*+]|\s*\d+\.)\s+", lines[i]):
                buf.append(lines[i])
                i += 1
            yield ("list", "\n".join(buf))
            continue
        # paragraph
        buf = [line]
        i += 1
        while i < n and lines[i].strip() and not lines[i].strip().startswith(("#", "|", "```", "-", "*", "+")) \
                and not re.match(r"^\d+\.\s+", lines[i]):
            buf.append(lines[i])
            i += 1
        yield ("para", "\n".join(buf))


def _split_table_row(row: str) -> list[str]:
    row = row.strip()
    if row.startswith("|"):
        row = row[1:]
    if row.endswith("|"):
        row = row[:-1]
    # naive split on unescaped pipes
    cells = []
    cur = []
    for ch in row:
        if ch == "|":
            cells.append("".join(cur).strip())
            cur = []
        else:
            cur.append(ch)
    cells.append("".join(cur).strip())
    return cells


def _md_inline_to_plain(text: str) -> str:
    """Strip bold / italics / inline code / links to plain text for docx cell."""
    text = re.sub(r"!\[[^\]]*\]\([^)]*\)", "[img]", text)  # image → placeholder
    text = re.sub(r"\[([^\]]*)\]\([^)]*\)", r"\1", text)   # link → label
    text = re.sub(r"`([^`]*)`", r"\1", text)               # inline code
    text = re.sub(r"\*\*([^*]*)\*\*", r"\1", text)         # bold
    text = re.sub(r"\*([^*]*)\*", r"\1", text)             # italic
    return text


def convert(md_path: Path, out_path: Path, title: str | None = None) -> int:
    try:
        from docx import Document
        from docx.shared import Pt
    except ImportError:
        print(
            "ERROR: python-docx 未安装。可选安装：pip install python-docx\n"
            "（离线导出依赖；飞书发布请用 prd_publish.py，不需此依赖）",
            file=sys.stderr,
        )
        return 2

    text = md_path.read_text(encoding="utf-8")
    text = _strip_frontmatter(text)

    doc = Document()
    if title:
        doc.add_heading(title, level=0)

    for kind, content in _iter_blocks(text):
        if kind == "h1":
            doc.add_heading(_md_inline_to_plain(content), level=1)
        elif kind == "h2":
            doc.add_heading(_md_inline_to_plain(content), level=2)
        elif kind == "h3":
            doc.add_heading(_md_inline_to_plain(content), level=3)
        elif kind == "code":
            p = doc.add_paragraph()
            run = p.add_run(content)
            run.font.name = "Consolas"
            run.font.size = Pt(9)
        elif kind == "table":
            rows = [_split_table_row(r) for r in content.splitlines()]
            if not rows:
                continue
            # drop separator row (---...)
            body_rows = [r for r in rows if not all(re.fullmatch(r"[\s:\-|]*", c or " ") for c in r)]
            if not body_rows:
                continue
            ncols = max(len(r) for r in body_rows)
            tbl = doc.add_table(rows=len(body_rows), cols=ncols)
            tbl.style = "Light Grid Accent 1"
            for ri, row in enumerate(body_rows):
                for ci in range(ncols):
                    cell_text = row[ci] if ci < len(row) else ""
                    tbl.cell(ri, ci).text = _md_inline_to_plain(cell_text)
        elif kind == "list":
            for line in content.splitlines():
                clean = re.sub(r"^(\s*)[-*+]\s+", r"\1• ", line)
                clean = re.sub(r"^(\s*)\d+\.\s+", r"\1", clean)
                doc.add_paragraph(_md_inline_to_plain(clean), style="List Bullet" if line.strip()[0] in "-*+" else "List Number")
        else:  # para
            doc.add_paragraph(_md_inline_to_plain(content))

    out_path.parent.mkdir(parents=True, exist_ok=True)
    doc.save(str(out_path))
    print(f"✅ 已导出 {out_path} ({out_path.stat().st_size} bytes)")
    return 0


def main() -> int:
    p = argparse.ArgumentParser(description=__doc__)
    p.add_argument("input", type=Path, help="Markdown PRD 或产物路径")
    p.add_argument("--output", type=Path, default=None, help="输出 .docx 路径（默认 <input>.docx）")
    p.add_argument("--title", default=None, help="文档标题（默认用文件名）")
    a = p.parse_args()

    if not a.input.is_file():
        print(f"ERROR: 输入文件不存在: {a.input}", file=sys.stderr)
        return 1
    out = a.output or a.input.with_suffix(".docx")
    title = a.title or a.input.stem
    return convert(a.input, out, title)


if __name__ == "__main__":
    sys.exit(main())
